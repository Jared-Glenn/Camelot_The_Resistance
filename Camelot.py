import pdb
import copy
import os
import random
import shutil
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Length, Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

# get_role_descriptions - this is called when information files are generated.
def get_role_description(role):
    return {
        # Good Roles
        'King Arthur' : 'You know each player that is seeking Excalibur in the right place. That is, each player that gains some benefit from Excalibur being in the location Excalibur is actually hidden.\n\nOBTAIN EXCALIBUR: (Must have declared as a Sword Seeker, King Arthur, or Sir Accolon [your choice].)\nAfter you declare, each time you are given the Scabbard, you may search the current location for Excalibur instead of changing the location. If it is in that location, you gain Excalibur and can use the Expend Excalibur ability.\n\nEXPEND EXCALIBUR: (Once per game. Must have found Excalibur.)\nWhen you attend a quest, once all of the quest cards are returned but before they are read, you may force the quest to succeed even when it would otherwise fail.\n',
        'Sir Bedivere' : 'SUSPEND: (You gain this ability when you gain the Scabbard, but choose not to move Excalibur)\nOnce per game, while Excalibur is in the Lake, after the quest cards have been collected for a quest, but before they are read, you can declare as a Guardian of Truth, Sir Bedivere, or Sir Agravaine (your choice). If you do, you may look at the quest cards before the leader and remove one of them. If you remove a Failure card, the next time you attend a quest, you MUST play a Failure card unless the Holy Grail is played on you.\n',
        'Sir Bertilak, The Green Knight' : 'You are a member of the Fairy Court and benefit from corrupting the Holy Grail.\n',
        'Sir Bors' : 'You have access to Cleanse cards.\n\nCLEANSE CARDS:\nIf there is a Rot or Regrowth card revealed on the same quest, negate all Cleanse, Rot, and Regrowth cards.\nIf no Rot or Regrowth cards are revealed on the same quest, Cleanse cards count as Failure cards.\n\nWhile Excalibur is at Camelot, you may choose to declare as Sir Bors. After you declare, the next time you successfully remove a Rot or Regrowth from a quest with a Cleanse card, you may Cleanse the Holy Grail as well, moving the Holy Grail one step further from Corruption.\n',
        'Sir Dagonet' : 'You cannot speak, but can communicate through gibberish sounds and body language.\nYou know King Arthur.\nYou appear Evil to Merlin and to all Evil players.\nOnly Sir Ector may know if Sir Dagonet is in this game.\n',
        'Sir Ector' : 'You know which Good roles are in the game, but not who has any given role.\n',
        'Sir Galahad' : 'REVEAL: (You gain this ability when you attend the First or Fourth Quest.)\nOnce per game, while Excalibur is at Camelot, you may declare as a Knight of Camelot, Sir Galahad, or Sir Lancelot (your choice). When you do, you instruct all players to close their eyes and hold their fists out in front of them. Name one good role (for example, King Arthur, Sir Gawain, etc.). If a player has that role, they must raise their thumb to indicate they are playing that role. You can then instruct all players to put their hands down, open their eyes, and resume play normally.',
        'Sir Gawain' : 'You know all members of the Fairy Court, Good and Evil.\nYour presence has caused the Grail to start closer to corruption.\n',
        'Queen Guinevere' : 'You know two \"rumors\" about other players.\nRumors tell you who knows something about another player, but does not explain that connection. You must use your wits to determine how the pieces fit together.\n',
        'Lady Iseult' : 'The person you see is also Good and is aware that you are Good.\nYou can be assassinated as either a Lover (which includes either Sir Tristan or Lady Iseult) or as Lady Iseult specifically.\nYou and Sir Tristan each know one location that Excalibur cannot be found.\n.',
        'Sir Lamorak' : 'You can see two pairs of players.\nOne pair of players are on opposite sides (Good and Evil or Sir Pelinor and the Questing Beast), and the other pair are on the same side (Evil and Evil or Good and Good).',
        'Merlin' : 'You know which people have Evil roles, but not who has any specific role.\n',
        'Lady Nimue, The Lady of the Lake' : 'You are a member of the Fairy Court and benefit from corrupting the Holy Grail.\n',
        'Sir Percival' : 'You know which people have the Merlin and Lady Morgana roles, but not who has each.\n',
        'Queen Titania of the Seelie Fairies' : 'You are a member of the Fairy Court and benefit from corrupting the Holy Grail.\nYou appear as Evil to all players with Evil roles (except Sir Colgrevance).\n',
        'Sir Tristan' : 'The person you see is also Good and is aware that you are Good.\nYou can be assassinated as either a Lover (which includes either Sir Tristan or Lady Iseult) or as Sir Tristan specifically.\nYou and Lady Iseult each know one location that Excalibur cannot be found.\n.',
        'King Uther' : 'EXILE: (You gain this ability when you vote against your own quest proposal or a quest proposal in which you have been chosen to attend.)\nOnce per game, while Excalibur is in the Stone, you may declare as a King of the Realm, King Uther, or King Vortigurn (your choice). You may only do this after a new leader is selected but before a quest vote occurs. If you do, you may select one player to be exiled from the game until the next quest is completed. That player is required to view your role information, and will see which role you possess. The exiled player must leave the play area to view this information, and you must be the one to go retrieve that player, affording you a moment of privacy with that player, if you wish.',

        # Evil Roles
        'Sir Accolon' : 'You know King Arthur.\n\nOBTAIN EXCALIBUR: (Once per game. Must have obtained the Scabbard twice [by any means].)\nOnce per game, when you obtain the Scabbard, you may search the current location for Excalibur instead of changing the location. If it is in that location, you gain Excalibur and can use the Expend Excalibur ability.\n\nEXPEND EXCALIBUR: (Once per game. Must have found Excalibur.)\nWhen you attend a quest, once all of the quest cards are returned but before they are read, you may force the quest to fail even when it would otherwise succeed.\n\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Sir Agravaine' : 'You have the Suspend ability.\n\nSUSPEND:\nOnce per game, while Excalibur is in the Lake, after the quest cards have been collected for a quest, but before they are read, you can declare as a Guardian of Truth, Sir Bedivere, or Sir Agravaine (your choice). If you do, you may look at the quest cards before the leader and remove one of them. If you remove a Failure card, the next time you attend a quest, you MUST play a Failure card unless the Holy Grail is played on you.\n\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Lady Annowre' : 'You know where Excalibur may be retrieved.\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Sir Colgrevance' : 'You know not only who else is Evil, but what role each other Evil player possesses.\nEvil players know that there is a Sir Colgrevance, but do not know that it is you or even that you are Evil.',
        'Sir Lancelot' : 'You have the Reveal ability.\n\nREVEAL:\nOnce per game, while Excalibur is at Camelot, you may declare as a Knight of Camelot, Sir Galahad, or Sir Lancelot (your choice). When you do, you instruct all players to close their eyes and hold their fists out in front of them. Name one good role (for example, King Arthur, Sir Gawain, etc.). If a player has that role, they must raise their thumb to indicate they are playing that role. You can then instruct all players to put their hands down, open their eyes, and resume play normally.\n\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Queen Mab of the Unseelie Fairies' : 'You are a member of the Fairy Court and benefit from corrupting the Holy Grail.\nYou ignore the effects of the Holy Grail.\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Sir Maleagant' : 'Each time you are given the Scabbard, you may declare as Sir Maleagant to force the next quest to have one additional knight attend it. This ability cannot be used on the final quest. This ability cannot be used if you use the Leader role to take the Scabbard.\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Sir Mordred' : 'You are hidden from all Good roles that could reveal that information.\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Lady Morgana' : 'You appear like Merlin to Sir Percival.\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'King Oberon of the Fairy Court' : 'You are a member of the Fairy Court and benefit from corrupting the Holy Grail.\nYou know Sir Gawain and Sir Bertilak, The Green Knight.\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',
        'Sir Palamedes' : 'You know Lady Iseult and you have the Exile ability.\n\nEXILE:\nWhile Excalibur is in the Stone, you may declare as a King of the Realm, King Uther, or King Vortigurn (your choice). You may only do this after a new leader is selected but before a quest vote occurs. If you do, you may select one player to be exiled from the game until the next quest is completed. That player is required to view your role information, and will see which role you possess. The exiled player must leave the play area to view this information, and you must be the one to go retrieve that player, affording you a moment of privacy with that player, if you wish.\n\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n\nATTENTION TRISTAN PLAYER:\nIf you have been Exiled by Sir Palamedes, you are now ONLY allowed to play Failure cards on quests unless affected by the Holy Grail (in which case, you must play a Success card).\nYou are still on the Good team and your Victory Points are unchanged.\nYou are still a valid Assassination Target.\nWhen Sir Palamedes retrieves you, you must reveal that you are Tristan to him.\n',
        'King Vortigurn' : 'EXILE:\nOnce per game, while Excalibur is in the Stone, you may declare as a King of the Realm, King Uther, or King Vortigurn (your choice). You may only do this after a new leader is selected but before a quest vote occurs. If you do, you may select one player to be exiled from the game until the next quest is completed. That player is required to view your role information, and will see which role you possess. The exiled player must leave the play area to view this information, and you must be the one to go retrieve that player, affording you a moment of privacy with that player, if you wish.\n\nLike other Evil characters, you know who else is Evil (except Sir Colgrevance).\n',

        # Neutral Roles
        'Sir Kay' : 'You are neutral and equally pulled to the Good and Evil sides, but you do have one ally who is either Good or Evil. You must determine if this ally is Good or Evil and assist as best you can.\nNiether Evil players nor Merlin can identify you as Good or Evil.',
        'Sir Pelinor' : 'You are Neutral in this battle and have no allies in this game.\nYour nemesis is The Questing Beast, who is also Neutral.\nYou have the Battle the Beast ability.\n\nBATTLE THE BEAST: (You must be on the Fifth Quest to use this ability)\nAfter all quest cards have been collected, but before they are read, you may declare as Sir Pelinor. The Questing Beast must reveal itself to all players.\n\nABOUT THE QUESTING BEAST:\nThe Questing Beast can see who you are.\nThe Questing Beast must play at least one \"The Questing Beast Was Here\" card or will lose significant Victory Points.\nIt may play a \"Reversal\" card once per game.\n',
        'The Questing Beast' : 'You are Neutral in this battle and have no allies in this game.\nYour nemesis is Sir Pelinor, who is also Neutral.\n\nABOUT PELINOR:\nSir Pelinor cannot see you, though you can see him.\nPelinor also wants to reach the Fifth Quest to gain additional Victory Points.\nBeware! If Sir Pelinor suspects you are on the Fifth Quest, he may declare as Sir Pelinor, causing you to reveal yourself and lose Victory Points.\n',

}.get(role,'ERROR: No description available.')

# get_role_information: this is called to populate information files
# blank roles:
# - Sir Lancelot: no information
# - King Arthur: no information
# - Queen Guinevere: too complicated to generate here
# - Sir Colgrevance: name, role (evil has an update later to inform them about the presence of Sir Colgrevance)
def get_role_information(my_player,players,relics):
    return {
        'Sir Tristan' : [[f'{relic.decoy2}' for relic in relics if relic.name == 'Excalibur'], ['{} is Lady Iseult.'.format(player.name) for player in players if player.role == 'Lady Iseult']],
        'Lady Iseult' : [[f'{relic.decoy1}' for relic in relics if relic.name == 'Excalibur'], ['{} is Sir Tristan.'.format(player.name) for player in players if player.role == 'Sir Tristan']],
        'Merlin' : ['{} is Evil'.format(player.name) for player in players if (player.team == 'Evil' and player.role != 'Sir Mordred') or player.role == 'Sir Dagonet'],
        'Sir Percival' : ['{} is Merlin or Lady Morgana.'.format(player.name) for player in players if player.role == 'Merlin' or player.role == 'Lady Morgana'],
        'King Arthur' : [f'{player.name} is seeking Excalibur in the correct location.' for relic in relics if relic.name == 'Excalibur' for role in relic.location_seeker for player in players if player.role == role],
        'Queen Titania of the Seelie Fairies' : [],
        'Lady Nimue, The Lady of the Lake' : [],
        'Sir Galahad' : [],
        'Queen Guinevere' : [str(get_rumors(my_player, players,relics))],
        'Sir Lamorak' : [str(get_relationships(my_player, players))],
        'Sir Ector' : [f'{player.role} is in this game.' for player in players if player.team == 'Good' and player.role != 'Sir Ector'],
        'Sir Dagonet' : ['{} is King Arthur.'.format(player.name) for player in players if player.role == 'King Arthur'],
        'King Uther' : [],
        'Sir Bedivere' : [],
        'Sir Gawain' : [f'{player.name} is a member of the Fairy Court.' for player in players if player.origin == 'Fae'],
        'Sir Bors' : [],
        'Sir Bertilak, The Green Knight' : [],

        'Sir Mordred' : ['{} is Evil.'.format(player.name) for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'Lady Morgana' : ['{} is Evil.'.format(player.name) for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'Sir Maleagant' : ['{} is Evil.'.format(player.name) for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'Sir Agravaine' : ['{} is Evil.'.format(player.name) for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'Sir Colgrevance' : ['{} is {}.'.format(player.name, player.role) for player in players if player.team == 'Evil' and player != my_player],
        'Sir Accolon' : [[f'{player.name} is King Arthur.' for player in players if player.role == 'King Arthur'], [f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')]],
        'Sir Lancelot' : [f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'King Vortigurn' : [f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'Lady Annowre' : [[f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')], [f'{relic.location}' for relic in relics if relic.name == 'Excalibur']],
        'Sir Palamedes' : [[f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')], [f'{player.name} is Lady Iseult.' for player in players if player.role == 'Lady Iseult']],
        'Queen Mab of the Unseelie Fairies' : [f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')],
        'King Oberon of the Fairy Court' : [[f'{player.name} is Evil.' for player in players if (player.team == 'Evil' and player != my_player and player.role != 'Sir Colgrevance') or (player.role == 'Queen Titania of the Seelie Fairies') or (player.role == 'Sir Dagonet')], [f'{player.name} is Sir Gawain.' for player in players if player.role == 'Sir Gawain'], [f'{player.name} is Sir Bertilak, The Green Knight.' for player in players if player.role == 'Sir Bertilak, The Green Knight']],

        'Sir Pelinor' : [],
        'The Questing Beast' : ['{} is Sir Pelinor.'.format(player.name) for player in players if player.role == 'Sir Pelinor'],
        'Sir Kay' : [str(get_ally(my_player, players))],
    }.get(my_player.role,[])

def get_role_victory_points(role):
    return {
        # Good Roles
        'King Arthur' : 'If you expended Excalibur\'s power, all other Good players gain 1 Victory Point and you gain 2 Victory Points.\n       If you do not obtain Excalibur, you lose 1 Victory Point.\n',
        'Sir Bedivere' : 'If you used your Suspend power, gain 1 Victory Point.\n',
        'Sir Bertilak, The Green Knight' : 'NONE',
        'Sir Bors' : 'For each time you used Cleanse to negate a Rot or Regrowth, gain 1 Victory Point.\n      For each time you used Cleanse and did not negate a Rot or Regrowth, lose 1 Victory Point.\n',
        'Sir Dagonet' : 'NONE',
        'Sir Ector' : 'If you are assassinated, you lose 3 Victory Points and all Evil players gain 3 Victory Points.\n',
        'Sir Galahad' : 'If you used your Reveal power, gain 1 Victory Point.\n',
        'Sir Gawain' : 'If you are assassinated, you lose 2 Victory Points, all Evil players gain 1 Victory Point, and all Fae players gain 4 Victory Points.\n',
        'Queen Guinevere' : 'If you are assassinated, you lose 2 Victory Points and all Evil players gain 2 Victory Points.\n',
        'Lady Iseult' : 'If you are assassinated as Lady Iseult, you lose 4 Victory Points and all Evil players gain 4 Victory Points.\n        If you are assassinated as a Lover (and not named as Lady Iseult), you lose 3 Victory Points and all Evil players gain 3 Victory Points.\n        If Sir Tristan is Assassinated, you lose 2 Victory Points.\n',
        'Sir Lamorak' : 'If you are assassinated, you lose 3 Victory Points and all Evil players gain 3 Victory Points.\n',
        'Merlin' : 'If you are assassinated, you lose 4 Victory Points and all Evil players gain 4 Victory Points.\n',
        'Lady Nimue, The Lady of the Lake' : 'If Excalibur is in the Lake at the end of the game, you gain 1 Victory Point.\n',
        'Sir Percival' : 'If Merlin is assassinated, you lose 3 Victory Points.\n',
        'Queen Titania of the Seelie Fairies' : 'NONE',
        'Sir Tristan' : 'If you are assassinated as Sir Tristan, you lose 4 Victory Points and all Evil players gain 4 Victory Points.\n        If you are assassinated as a Lover (and not named as Sir Tristan), you lose 3 Victory Points and all Evil players gain 3 Victory Points.\n        If Lady Iseult is Assassinated, you lose 2 Victory Points.\n',
        'King Uther' : 'If you used your Exile power, gain 1 Victory Point.\n',

        # Evil Roles
        'Sir Accolon' : 'If you expended Excalibur\'s power, all other Evil players gain 1 Victory Point and you gain 2 Victory Points.\n       If you do not obtain Excalibur, you lose 1 Victory Point.',
        'Sir Agravaine' : 'NONE',
        'Lady Annowre' : 'If are in possession of the Excalibur card for the correct location at the end of the game, you gain 3 Victory Points.\n        Each time you are caught taking an Excalibur card, you must return it to the Excalibur stack and you lose 1 Victory Point.\n        If Excalibur\'s power is expended, you lose 5 Victory Points.\n',
        'Sir Colgrevance' : 'NONE',
        'Sir Lancelot' : 'If the assassination attempt is successful at all, you lose 3 Victory Points.\n          If YOU are assassinated (while a valid target), you lose an additional 3 Victory Points.\n          If the Assassination Attempt fails and exactly two quests have failed, you gain 3 Victory Points.\n',
        'Queen Mab of the Unseelie Fairies' : 'NONE',
        'Sir Maleagant' : 'Each time the Scabbard is given to you, you gain 1 Victory Point.\n           If the Scabbard is never given to you, you lose 2 Victory Points.\n',
        'Sir Mordred' : 'At the end of the game, before the Assassination Attempt, choose one player. If that player loses one or more Victory Points, you may gain the same number of Victory Points.\n         If your chosen player does not lose any Victory Points, that player gains 2 Victory Points and you lose 2 Victory Points.\n',
        'Lady Morgana' : 'NONE',
        'King Oberon of the Fairy Court' : 'NONE',
        'Sir Palamedes' : 'If you Exiled Sir Tristan, you gain 3 Victory Points.\n           If you Exiled a player that is not Sir Tristan, you lose 1 Victory Point.\n           If Sir Tristan or Lady Iseult are assassinated, you lose 6 Victory Points.\n',
        'King Vortigurn' : 'NONE',

        # Neutral Roles
        'Sir Kay' : 'Your Victory Point total is equal to the Victory Point total of your ally.\n',
        'Sir Pelinor' : 'If you attended the Fifth Quest, you gain 2 Victory Points.\n         If you declared as Sir Pelinor while the Questing Beast was on the Fifth Quest with you, you gain 3 Victory Points.\n         If you declared as Sir Pelinor while the Questing Beast was not on the Fifth Quest with you, you lose 5 Victory Points.\n         If no The Questing Beast Was Here cards were played before the Fifth Quest, you gain 5 Victory Points.\n',
        'The Questing Beast' : 'If you DID attend the Fifth Quest and Sir Pelinor did NOT declare, you gain 5 Victory Points.\n                    If you DID attend the Fifth Quest and Sir Pelinor DID declare, you lose 3 Victory Points.\n                    If you did NOT attend the Fifth Quest and Sir Pelinor DID declare, you gain 5 Victory Points.\n                    If you did NOT attend the Fifth Quest and Sir Pelinor did NOT declare, you lose 3 Victory Points.\n                    If no The Questing Beast Was Here cards were played before the Fifth Quest, you lose 5 Victory Points.\n',
}.get(role,'ERROR: No description available.')

def get_playable_cards(role):
    return {
        # Good Roles
        'King Arthur' : 'Success\n',
        'Sir Bedivere' : 'Success\n',
        'Sir Bertilak, The Green Knight' : 'Regrowth\n',
        'Sir Bors' : 'Success\nCleanse\n',
        'Sir Dagonet' : 'Success\n',
        'Sir Ector' : 'Success\n',
        'Sir Galahad' : 'Success\n',
        'Sir Gawain' : 'Success\n',
        'Queen Guinevere' : 'Success\n',
        'Lady Iseult' : 'Success\n',
        'Sir Lamorak' : 'Success\n',
        'Merlin' : 'Success\n',
        'Lady Nimue, The Lady of the Lake' : 'Regrowth\n',
        'Sir Percival' : 'Success\n',
        'Queen Titania of the Seelie Fairies' : 'Regrowth\n',
        'Sir Tristan' : 'Success\n',
        'King Uther' : 'Success\n',

        # Evil Roles
        'Sir Accolon' : 'Success\nFailure\n',
        'Sir Agravaine' : 'Success\nFailure\n',
        'Lady Annowre' : 'Success\nFailure\n',
        'Sir Colgrevance' : 'Success\nFailure\n',
        'Sir Lancelot' : 'Success\nFailure\n',
        'Queen Mab of the Unseelie Fairies' : 'Rot\n',
        'Sir Maleagant' : 'Failure\n',
        'Sir Mordred' : 'Success\nFailure\n',
        'Lady Morgana' : 'Success\nFailure\n',
        'King Oberon of the Fairy Court' : 'Rot\nRegrowth\n',
        'Sir Palamedes' : 'Success\nFailure\n',
        'King Vortigurn' : 'Success\nFailure\n',

        # Neutral Roles
        'Sir Kay' : 'Success\nFailure\nRot\nRegrowth\n',
        'Sir Pelinor' : 'Success\nReversal\n',
        'The Questing Beast' : 'The Questing Beast Was Here\n',
}.get(role,'ERROR: No description available.')

def get_conditional_cards(role):
    return {
        # Good Roles
        '1-King Arthur' : 'NONE',
        '3-King Arthur' : 'NONE',
        '1-Sir Bedivere' : 'NONE',
        '3-Sir Bedivere' : '[The Suspended Card] ',
        '4-Sir Bedivere' : '(Must be played the next time you attend a quest.)\n',
        '1-Sir Bertilak, The Green Knight' : 'Rot ',
        '2-Sir Bertilak, The Green Knight' : '(Must be played when the Holy Grail is played on you.)\n',
        '3-Sir Bertilak, The Green Knight' : '[Any Card] ',
        '4-Sir Bertilak, The Green Knight' : '(Available for each quest after the Holy Grail has been corrupted.)\n',
        '1-Sir Bors' : 'NONE',
        '3-Sir Bors' : 'NONE',
        '1-Sir Dagonet' : 'NONE',
        '3-Sir Dagonet' : 'NONE',
        '1-Sir Ector' : 'NONE',
        '3-Sir Ector' : 'NONE',
        '1-Sir Galahad' : 'NONE',
        '3-Sir Galahad' : 'NONE',
        '1-Sir Gawain' : 'Reversal ',
        '2-Sir Gawain' : '(If there are no Fae players in this game.)\n',
        '3-Sir Gawain' : 'NONE',
        '1-Queen Guinevere' : 'NONE',
        '1-Lady Iseult' : 'NONE',
        '1-Sir Lamorak' : 'NONE',
        '1-Merlin' : 'Reversal ',
        '2-Merlin' : '(Available when Excalibur is in the Lake, unless the Holy Grail is currently played on you.)\n',
        '3-Merlin' : 'NONE',
        '1-Lady Nimue, The Lady of the Lake' : 'Success ',
        '2-Lady Nimue, The Lady of the Lake' : '(Must be played when the Holy Grail is played on you, unless Excalibur is in the Lake.)\n',
        '3-Lady Nimue, The Lady of the Lake' : '[Any Card] ',
        '4-Lady Nimue, The Lady of the Lake' : '(Available for each quest after the Holy Grail has been corrupted.)\n',
        '1-Sir Percival' : 'Reversal ',
        '2-Sir Percival' : '(Available when Excalibur is at Camelot, unless the Holy Grail is currently played on you.)\n',
        '3-Sir Percival' : 'NONE',
        '1-Queen Titania of the Seelie Fairies' : 'NONE',
        '3-Queen Titania of the Seelie Fairies' : '[Any Card] ',
        '4-Queen Titania of the Seelie Fairies' : '(Available for each quest after the Holy Grail has been corrupted.)\n',
        '1-Sir Tristan' : 'NONE',
        '1-King Uther' : 'NONE',

        # Evil Roles
        '1-Sir Accolon' : 'NONE',
        '1-Sir Agravaine' : 'NONE',
        '1-Lady Annowre' : 'NONE',
        '1-Sir Colgrevance' : 'NONE',
        '1-Sir Lancelot' : 'NONE',
        '1-Queen Mab of the Unseelie Fairies' : 'NONE',
        '3-Queen Mab of the Unseelie Fairies' : '[Any Card] ',
        '4-Queen Mab of the Unseelie Fairies' : '(Available for each quest after the Holy Grail has been corrupted.)\n',
        '1-Sir Maleagant' : 'NONE',
        '1-Sir Mordred' : 'NONE',
        '1-Lady Morgana' : 'Reversal ',
        '2-Lady Morgana' : '(Available when Excalibur is in the Stone, unless the Holy Grail is currently played on you.)\n',
        '3-Lady Morgana' : 'NONE',
        '1-King Oberon of the Fairy Court' : 'NONE',
        '3-King Oberon of the Fairy Court' : '[Any Card] ',
        '4-King Oberon of the Fairy Court' : '(Available for each quest after the Holy Grail has been corrupted.)\n',
        '1-Sir Palamedes' : 'Success ',
        '2-Sir Palamedes' : '(Must be played on each quest after you have Exiled Sir Tristan.)\n',
        '3-Sir Palamedes' : 'NONE',
        '1-King Vortigurn' : 'NONE',

        # Neutral Roles
        '1-Sir Kay' : 'Success ',
        '2-Sir Kay' : '(Must be played when the Holy Grail is played on you.)\n',
        '3-Sir Kay' : 'NONE',
        '1-Sir Pelinor' : 'NONE',
        '1-The Questing Beast' : 'Reversal ',
        '2-The Questing Beast' : '(Available once per game.)\n',
        '3-The Questing Beast' : 'NONE'
}.get(role,'NONE')

def get_rumors(my_player, players, relics):
    rumors = []

    if my_player.role != 'Queen Guinevere':
        return
    
    # Generate rumors about Merlin
    merlin_player = None
    is_Merlin = 0
    for player in players:
        if player.role == 'Merlin':
            merlin_player = player.name
            is_Merlin = 1
    if is_Merlin == 1:
        for player in players:
            if (player.team == 'Evil' and player.role != 'Sir Mordred') or player.role == "Sir Dagonet":
                player_of_evil = player.name
                rumors.append('{} sees {}'.format(merlin_player, player_of_evil))

    # Generate rumors about Sir Percival
    percival_player = None
    is_Percival = 0
    for player in players:
        if player.role == 'Sir Percival':
            percival_player = player.name
            is_Percival = 1
    if is_Percival == 1:
        for player in players:
            if player.role == 'Merlin' or player.role == 'Lady Morgana':
                seer = player.name
                rumors.append('{} sees {}'.format(percival_player, seer))

    # Generate rumor about the Lovers
    tristan_player = None
    iseult_player = None
    is_Lovers = 0
    for player in players:
        if player.role == 'Sir Tristan':
            tristan_player = player.name
            is_Lovers += 1
        elif player.role == 'Lady Iseult':
            iseult_player = player.name
            is_Lovers += 1
    if is_Lovers == 2:
        rumors.append('{} sees {}'.format(tristan_player, iseult_player))
        rumors.append('{} sees {}'.format(iseult_player, tristan_player))

    # Generate rumor about Evil players
    for player in players:
        if player.team == 'Evil' and player.role != 'Sir Mordred':
            for player_two in players:
                if (player_two.team == 'Evil' and player_two.role != 'Sir Mordred' and player_two.role != 'Sir Colgrevance' and player_two != player) or (player_two.role == 'Queen Titania of the Seelie Fairies' and player_two != player):
                    rumors.append('{} sees {}'.format(player.name, player_two.name))

    # Generate rumor about Sir Ector
    is_Ector = 0
    for player in players:
        if player.role == 'Sir Ector':
            is_Ector = 1
    if is_Ector == 1:
        for player in players:
            if player.team == 'Good' and player.role != 'Sir Ector' and player.role != 'Queen Guinevere':
                rumors.append(f'Sir Sir Ector sees {player.role}')

    # Generate rumor about The Questing Beast
    questing_player = None
    is_Questing = 0
    for player in players:
        if player.role == 'The Questing Beast':
            questing_player = player
            is_Questing = 1
    if is_Questing == 1:
        for player in players:
            if player.role == 'Sir Pelinor':
                rumors.append(f'{questing_player} sees {player}.')

    # Generate rumors about King Arthur
    arthur_player = None
    is_Arthur = 0
    for player in players:
        if player.role == 'King Arthur':
            is_Arthur = 1
            arthur_player = player
    for relic in relics:
        if relic == 'Excalibur':
            for role in relic.location_seeker:
                for player in players:
                    if player.role == role:
                        rumors.append(f'{arthur_player.name} sees {player.name}')

    # Pick two rumors and return them as a string.
    rumor_one = random.choice(rumors)
    rumor_two = random.choice(rumors)
    while rumor_one == rumor_two:
            rumor_two = random.choice(rumors)
    return rumor_one + '\n' + rumor_two

def get_relationships(my_player, players):

    if my_player.role != 'Sir Lamorak':
        return
    
    # Assign teams
    good_team = []
    evil_team = []
    neutral_team = []
    for player in players:
        if player.team == 'Good' and player.role != 'Sir Lamorak':
            good_team.append(player)
        if player.team == 'Evil' and player.role != 'Sir Mordred':
            evil_team.append(player)
        if player.team == 'Neutral' and player.role != 'Sir Kay':
            neutral_team.append(player)
    valid_players = good_team + evil_team + neutral_team

    # Choose random Opposing Team players
    opposition = None
    opposing_player = random.choice(valid_players)
    if opposing_player.team == 'Good':
            opposition = opposing_player.name + ' opposes ' + (random.choice(evil_team)).name
    elif opposing_player.team == 'Evil':
            opposition = opposing_player.name + ' opposes ' + (random.choice(good_team)).name
    elif opposing_player.role == 'Sir Pelinor':
            for player in players:
                if player.role == 'The Questing Beast':
                    opposition = opposing_player.name + ' opposes ' + player.name
    elif opposing_player.role == 'The Questing Beast':
            for player in players:
                if player.role == 'Sir Pelinor':
                    opposition = opposing_player.name + ' opposes ' + player.name

    # Choose random Collaborator players
    # Random choice of good or evil team (else good would be much more likely)
    # while loop to prevent collaborators from being the same player
    collaboration = None
    player_one = "1"
    player_two = "1"
    random_team = random.choice(['Good','Evil'])
    while player_one == player_two:
        if random_team == 'Good':
            player_one = (random.choice(good_team)).name
            player_two = (random.choice(good_team)).name
        elif random_team == 'Evil':
            player_one = (random.choice(evil_team)).name
            player_two = (random.choice(evil_team)).name

    collaboration = player_one + ' is collaborating with ' + player_two
    return opposition + '\n' + collaboration


def get_ally(my_player, players):
    
    if my_player.role != 'Sir Kay':
        return
                    
    # Get Sir Kay's team and make a list of players on that team.
    allies = []
    if my_player.secret == 'Good':
        for player in players:
            if player.team == 'Good':
                allies.append(player)
    if my_player.secret == 'Evil':
        for player in players:
            if player.team == 'Evil':
                allies.append(player)

    # Return a random ally.
    kay_ally = random.choice(allies)
    return f'{kay_ally.name} is your ally. If {kay_ally.name} wins the game, so do you.'


# EXCALIBUR
# Randomly choose a location for Excalibur and keep track of decoy locations.
def get_excalibur():
    excalibur_hiding_places = [' in the Stone', ' at Camelot', ' in the Lake']

    # Randomly select one location for Excalibur to be
    excalibur_location = random.choice(excalibur_hiding_places)
    excalibur_hiding_places.remove(excalibur_location)

    # Produce the two decoy locations.
    excalibur_decoy1 = random.choice(excalibur_hiding_places)
    excalibur_hiding_places.remove(excalibur_decoy1)
    excalibur_decoy2 = random.choice(excalibur_hiding_places)

    return excalibur_location, excalibur_decoy1, excalibur_decoy2
          

class Player():
    # Players have the following traits
    # name: the name of the player as fed into system arguments
    # role: the role the player possesses
    # team: whether the player is on good, evil, or neutral's team
    # type: information or ability
    # seen: a list of what they will see
    # modifier: the random modifier this player has [NOT CURRENTLY UTILIZED]
    def __init__(self, name):
        self.name = name
        self.role = None
        self.team = None
        self.origin = None
        self.modifier = None
        self.info = []
        self.is_assassin = False
        self.secret = None

    def set_role(self, role):
        self.role = role

    def set_team(self, team):
        self.team = team
        
    def set_origin(self, origin):
        self.origin = origin

    def add_info(self, info):
        self.info += info

    def erase_info(self, info):
        self.info = []

    def generate_info(self, players):
        pass

class Relic():
    # Relics have the following traits
    # name: the name of the relic as fed into system arguments
    # type: the type the relic possesses
    # location: where the relic may be claimed
    # decoy1: first location where the relic is not found
    # decoy2: second location where the relic is not found
    # location_seeker: list of roles that seek the same location as the relic
    # decoy1_seeker: list of roles that seek the same location as decoy1
    # decoy2_seeker:list of roles that seek the same location as decoy2
    def __init__(self, name):
        self.name = name
        self.role = None
        self.location = None
        self.decoy1 = None
        self.decoy2 = None
        self.location_seeker = []
        self.decoy1_seeker = []
        self.decoy2_seeker = []

    def set_type(self, type):
        self.type = type

    def set_seeker(self, location):
        if location == ' in the Stone':
            return 'King Uther', 'King Vortigurn', 'Lady Morgana', 'Sir Palamedes', 'Sebile'
        elif location == ' at Camelot':
            return 'Sir Percival', 'Sir Galahad', 'Sir Lancelot'
        elif location == ' in the Lake':
            return 'Merlin', 'Sir Bedivere', 'Sir Agravaine', 'Lady Nimue, The Lady of the Lake'

    def set_location(self, location):
        seekers = self.set_seeker(location)
        self.location = 'Excalibur is' + location
        for seeker in seekers:
            self.location_seeker.append(seeker)

    def set_decoy1(self, decoy1):
        seekers = self.set_seeker(decoy1)
        self.decoy1 = 'Excalibur is not' + decoy1
        for seeker in seekers:
            self.decoy1_seeker.append(seeker)

    def set_decoy2(self, decoy2):
        seekers = self.set_seeker(decoy2)
        self.decoy2 = 'Excalibur is not' + decoy2
        for seeker in seekers:
            self.decoy2_seeker.append(seeker)

def get_player_info(player_names):
    num_players = len(player_names)
    if len(player_names) != num_players:
        print('ERROR: Duplicate player names.')
        exit(1)

    # Place Excalibur and decoy locations.
    relics = []
    excalibur_info = get_excalibur()
    excalibur = Relic("Excalibur")
    relics.append(excalibur)
    excalibur.set_type('Sword')
    excalibur.set_location(excalibur_info[0])
    excalibur.set_decoy1(excalibur_info[1])
    excalibur.set_decoy2(excalibur_info[2])

    # create player objects
    players = []
    for i in range(0, len(player_names)):
        player = Player(player_names[i])
        players.append(player)

    # number of good and evil roles
    num_neutral = 0
    if num_players <= 6:
        if random.choice([True, False, False]):
            kay_team = random.choice(['Good','Evil'])
            if kay_team == 'Good':
                num_evil = 2
                num_neutral = 1
            elif kay_team == 'Evil':
                num_evil = 1
                num_neutral = 1
        else:
            num_evil = 2
    elif num_players <= 8:
        if random.choice([True, False, False, False]):
            kay_team = random.choice(['Good','Evil'])
            if kay_team == 'Good':
                num_evil = 3
                num_neutral = 1
            elif kay_team == 'Evil':
                num_evil = 2
                num_neutral = 1
        else:
            num_evil = 3
    elif num_players <= 10:
        choice = random.choice(['kay', 'hunt', 'kayandhunt', 'niether', 'niether'])
        if choice == 'hunt':
            num_evil = 3
            num_neutral = 2
        elif choice == 'kay':
            kay_team = random.choice(['Good','Evil'])
            if kay_team == 'Good':
                num_evil = 4
                num_neutral = 1
            elif kay_team == 'Evil':
                num_evil = 3
                num_neutral = 1
        elif choice == 'kayandhunt':
            kay_team = random.choice(['Good','Evil'])
            if kay_team == 'Good':
                num_evil = 3
                num_neutral = 3
            elif kay_team == 'Evil':
                num_evil = 2
                num_neutral = 3
        elif choice == 'niether':
            num_evil = 4
    elif num_players <= 12:
        choice = random.choice(['kay', 'hunt', 'kayandhunt', 'niether', 'niether'])
        if choice == 'hunt':
            num_evil = 4
            num_neutral = 2
        elif choice == 'kay':
            kay_team = random.choice(['Good','Evil'])
            if kay_team == 'Good':
                num_evil = 5
                num_neutral = 1
            elif kay_team == 'Evil':
                num_evil = 4
                num_neutral = 1
        elif choice == 'kayandhunt':
            kay_team = random.choice(['Good','Evil'])
            if kay_team == 'Good':
                num_evil = 4
                num_neutral = 3
            elif kay_team == 'Evil':
                num_evil = 3
                num_neutral = 3
        elif choice == 'niether':
            num_evil = 5
    num_good = num_players - num_evil - num_neutral

    # establish available roles
    good_roles = ['Merlin', 'Sir Percival', 'King Arthur', 'Sir Dagonet']
    evil_roles = ['Sir Mordred', 'Lady Morgana', 'Sir Accolon', 'Lady Annowre']
    neutral_roles = ['Sir Kay']
                    
    # 6 plus
    if num_players > 5:
        good_roles.extend(['Sir Tristan', 'Lady Iseult', 'Queen Guinevere'])
        evil_roles.extend(['Sir Palamedes', 'Sir Maleagant'])

    # 7 plus
    if num_players > 6:
        good_roles.extend(['King Uther', 'Sir Bedivere', 'Sir Galahad', 'Sir Ector'])
        evil_roles.extend(['King Vortigurn', 'Sir Lancelot', 'Sir Agravaine'])

    # 8 plus
    if num_players > 7:
        good_roles.extend(['Lady Nimue, The Lady of the Lake', 'Sir Bertilak, The Green Knight', 'Sir Gawain', 'Sir Bors', 'Queen Titania of the Seelie Fairies'])
        evil_roles.extend(['Queen Mab of the Unseelie Fairies', 'King Oberon of the Fairy Court'])

    # 10 plus
    if num_players > 9:
        good_roles.append('Sir Lamorak')
        evil_roles.append('Sir Colgrevance')
        neutral_roles.extend(['Sir Pelinor', 'The Questing Beast'])

    good_roles_in_game = random.sample(good_roles, num_good)
    evil_roles_in_game = random.sample(evil_roles, num_evil)
    neutral_roles_in_game = []
    if num_neutral == 1:
        neutral_roles_in_game.append('Sir Kay')
    elif num_neutral == 2:
        neutral_roles_in_game.extend(['Sir Pelinor', 'The Questing Beast'])
    elif num_neutral == 3:
        neutral_roles_in_game.extend(['Sir Kay', 'Sir Pelinor', 'The Questing Beast'])
        

    # lone lovers are rerolled
    # 50% chance to reroll one lone lover
    # 50% chance to reroll another role into a lover
    if sum(gr in ['Sir Tristan','Lady Iseult'] for gr in good_roles_in_game) == 1 and num_good > 1:
        if 'Sir Tristan' in good_roles_in_game:
            good_roles_in_game.remove('Sir Tristan')
        if 'Lady Iseult' in good_roles_in_game:
            good_roles_in_game.remove('Lady Iseult')

        if random.choice([True, False]):
            # replacing the lone lover
            available_roles = good_roles
            not_available = good_roles_in_game + ['Sir Tristan'] + ['Lady Iseult']
            for role in not_available:
                available_roles.remove(role)
            # DecrecationWarning issue. Found solution at https://stackoverflow.com/questions/70426576/get-random-number-from-set-deprecation
            good_roles_in_game.append(random.choice(available_roles))
        else:
            # upgradng to pair of lovers
            rerolled = random.choice(good_roles_in_game)
            good_roles_in_game.remove(rerolled)
            good_roles_in_game.append('Sir Tristan')
            good_roles_in_game.append('Lady Iseult')

    # Add Sir Palamedes to the game. 50% chance if Sir Tristan and Lady Iseult are in the game.
    is_Palamedes = False
    if (sum(gr in ['Sir Tristan','Lady Iseult'] for gr in good_roles_in_game) == 2) and (num_good > 2):
        for er in evil_roles_in_game:
            if er == 'Sir Palamedes':
                is_Palamedes = True
                break
            else:
                continue
        if is_Palamedes == False and random.choice([True,False]):
            rerolled = random.choice(evil_roles_in_game)
            evil_roles_in_game.remove(rerolled)
            evil_roles_in_game.append('Sir Palamedes')
    
    # Prevent Sir Palamedes from being added to a game if there is no Sir Tristan and Lady Iseult
    if (sum(gr in ['Sir Tristan','Lady Iseult'] for gr in good_roles_in_game) == 0):
        available_roles = evil_roles
        for er in evil_roles_in_game:
            available_roles.remove(er)
            if er == 'Sir Palamedes':
                is_Palamedes = True
        if is_Palamedes == True:
            evil_roles_in_game.remove('Sir Palamedes')
            evil_roles_in_game.append(random.choice(available_roles))

    # Ensure there is a role that can be the Assassin.
    if num_evil == 1 and (sum(er in ['Sir Lancelot','Sir Palamedes', 'King Oberon of the Fairy Court'] for er in evil_roles_in_game) >= 1):
        incompatable_removed = 0
        for roles in evil_roles_in_game:
            if incompatable_removed == 0 and (role == 'Sir Lancelot' or role == 'Sir Palamedes' or role == 'King Oberon of the Fairy Court'):
                evil_roles_in_game.remove(role)
                incompatable_removed += 1
        available_roles = evil_roles
        not_available = evil_roles_in_game + ['Sir Lancelot'] + ['Sir Palamedes'] + ['King Oberon of the Fairy Court']
        for role in not_available:
            available_roles.remove(role)
        evil_roles_in_game.append(random.choice(available_roles))
    elif num_evil == 2 and (sum(er in ['Sir Lancelot','Sir Palamedes', 'King Oberon of the Fairy Court'] for er in evil_roles_in_game) >= 2):
        incompatable_removed = 0
        for roles in evil_roles_in_game:
            if incompatable_removed == 0 and (role == 'Sir Lancelot' or role == 'Sir Palamedes' or role == 'King Oberon of the Fairy Court'):
                evil_roles_in_game.remove(role)
                incompatable_removed += 1
        available_roles = evil_roles
        not_available = evil_roles_in_game + ['Sir Lancelot'] + ['Sir Palamedes'] + ['King Oberon of the Fairy Court']
        for role in not_available:
            available_roles.remove(role)
        evil_roles_in_game.append(random.choice(available_roles))
    elif num_evil == 3 and (sum(er in ['Sir Lancelot','Sir Palamedes', 'King Oberon of the Fairy Court'] for er in evil_roles_in_game) >= 3):
        incompatable_removed = 0
        for roles in evil_roles_in_game:
            if incompatable_removed == 0 and (role == 'Sir Lancelot' or role == 'Sir Palamedes' or role == 'King Oberon of the Fairy Court'):
                evil_roles_in_game.remove(role)
                incompatable_removed += 1
        available_roles = evil_roles
        not_available = evil_roles_in_game + ['Sir Lancelot'] + ['Sir Palamedes'] + ['King Oberon of the Fairy Court']
        for role in not_available:
            available_roles.remove(role)
        evil_roles_in_game.append(random.choice(available_roles))
            
            
    # roles after validation
    #print(good_roles_in_game)
    #print(evil_roles_in_game)

    # role assignment
    random.shuffle(players)
    neutral_players = []
    good_players = players[:num_good]
    if num_neutral == 0:
        evil_players = players[num_good:]
    else:
        goodAndEvil = num_good + num_evil
        evil_players = players[num_good:goodAndEvil]
        neutral_players = (players[goodAndEvil:])

    player_of_role = dict()

    # Good Team: Set Team and Origin.
    for gp in good_players:
        new_role = good_roles_in_game.pop()
        gp.set_role(new_role)
        gp.set_team('Good')
        player_of_role[new_role] = gp
        if gp.role == 'Sir Bertilak, The Green Knight' or gp.role == 'Lady Nimue, The Lady of the Lake' or gp.role == 'Queen Titania of the Seelie Fairies':
            gp.set_origin('Fae')
        else:
            gp.set_origin('Mortal')

    # Assign Compatable Evil Player to Assassin role.
    if evil_players[0].role != 'Sir Lancelot' and evil_players[0].role != 'Sir Palamedes' and evil_players[0].role != 'King Oberon of the Fairy Court':
        evil_players[0].is_assassin = True
    elif evil_players[1].role != 'Sir Lancelot' and evil_players[0].role != 'Sir Palamedes' and evil_players[0].role != 'King Oberon of the Fairy Court':
        evil_players[1].is_assassin = True
    elif evil_players[2].role != 'Sir Lancelot' and evil_players[0].role != 'Sir Palamedes' and evil_players[0].role != 'King Oberon of the Fairy Court':
        evil_players[2].is_assassin = True
    else:
        evil_players[3].is_assassin = True

    # Evil Team: Set Team and Origin.
    for ep in evil_players:
        new_role = evil_roles_in_game.pop()
        ep.set_role(new_role)
        ep.set_team('Evil')
        player_of_role[new_role] = ep
        if ep.role == 'Queen Mab of the Unseelie Fairies' or ep.role == 'King Oberon of the Fairy Court':
            ep.set_origin('Fae')
        else:
            ep.set_origin('Mortal')

    # Neutral Team: Set Team and Origin.
    for np in neutral_players:
        new_role = neutral_roles_in_game.pop()
        np.set_role(new_role)
        np.set_team('Neutral')
        player_of_role[new_role] = np
        np.set_origin('Mortal')
                    
    for player in players:
        if player.role == 'Sir Kay':
            player.secret = kay_team
            kay_name = player.name
            kay_player = player
    
    for p in players:
        p.add_info(get_role_information(p,players,relics))
        try:
            if isinstance(p.info[0], list):
                try:
                    p.info = p.info[0] + p.info[1]
                except Exception:
                    pass
        except Exception:
            pass
        random.shuffle(p.info)
        # print(p.name,p.role,p.team,p.info)

    # Informing Evil about Sir Colgrevance
    for ep in evil_players:
        if ep.role != 'Sir Colgrevance' and player_of_role.get('Sir Colgrevance'):
            ep.add_info(['Sir Colgrevance lurks in the shadows. (There is another Evil that you do not see.)'])
        if ep.role != 'Sir Colgrevance' and player_of_role.get('Queen Titania of the Seelie Fairies'):
            ep.add_info(['Queen Titania of the Seelie Fairies has infiltrated your ranks. (One of the people you see is not Evil.)'])

    # delete and recreate game directory
    if os.path.isdir("game"):
        shutil.rmtree("game")
    os.mkdir("game")

    for player in players:
        document = Document()
        sections = document.sections
        section = sections[0]
    
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
        run = paragraph.add_run("    " + player.name + ",")
        font = run.font
        font.name = 'Breathe Fire III'
        font.size = Pt(50)
    
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
        run = paragraph.add_run("you are ")
        font = run.font
        font.name = 'Breathe Fire III'
        font.size = Pt(30)
    
    
        if player.team == 'Good':
            run = paragraph.add_run(player.role+"\n")
            font = run.font
            font.name = 'Breathe Fire III'
            font.size = Pt(30)
            font.color.rgb = RGBColor(42, 96, 153) # Deep Blue
        elif player.team == 'Evil':
            run = paragraph.add_run(player.role+"\n")
            font = run.font
            font.name = 'Breathe Fire III'
            font.size = Pt(30)
            font.color.rgb = RGBColor(141, 40, 30) # Deep Red
        elif player.team == 'Neutral':
            run = paragraph.add_run(player.role+"\n")
            font = run.font
            font.name = 'Breathe Fire III'
            font.size = Pt(30)
            font.color.rgb = RGBColor(232, 162, 2) # Gold
    
        insertHR(paragraph)
    
        # Allegiance
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
        run = paragraph.add_run("\nAllegiance: ")
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(14)
        font.bold = True
    
        if player.team == 'Good':
            run = paragraph.add_run("Good")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.color.rgb = RGBColor(42, 96, 153) # Deep Blue
        if player.team == 'Evil':
            run = paragraph.add_run("Evil")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.color.rgb = RGBColor(141, 40, 30) # Deep Red
        if player.team == 'Neutral':
            run = paragraph.add_run("Neutral")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.color.rgb = RGBColor(232, 162, 2) # Gold
    
    
        # Origins
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
        run = paragraph.add_run("Origins: ")
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(14)
        font.bold = True
    
        if player.origin == 'Mortal':
            run = paragraph.add_run("Mortal")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.color.rgb = RGBColor(120, 3, 115) # Purple
        if player.origin == 'Fae':
            run = paragraph.add_run("Fae")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.color.rgb = RGBColor(30, 106, 57) # Green
    
        # Assassination
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
        if player.role == 'Sir Lancelot':
            run = paragraph.add_run("Assassination? ")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
            font.bold = True
            run = paragraph.add_run("CONDITIONAL (You may be assassinated if FEWER than two quests fail.)\n ")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
        elif player.team == 'Evil' and player.is_assassin == False:
            run = paragraph.add_run("Assassination? ")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
            font.bold = True
            run = paragraph.add_run("You are NOT the Assassin.\n")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
        elif player.is_assassin == True:
            run = paragraph.add_run("Assassination? ")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
            font.bold = True
            run = paragraph.add_run("You are the ASSASSIN.\n")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
            font.color.rgb = RGBColor(255, 0, 0) # Blood Red
        elif player.role == 'Sir Tristan' or player.role == 'Lady Iseult' or player.role == 'Merlin' or player.role == 'Queen Guinevere' or player.role == 'Sir Lamorak' or player.role == 'Sir Gawain' or player.role == 'Sir Ector':
            run = paragraph.add_run("Assassination? ")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
            font.bold = True
            run = paragraph.add_run("You are a TARGET for the Assassin.\n")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
        else:
            run = paragraph.add_run("Assassination? ")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True
            font.bold = True
            run = paragraph.add_run("You are NOT a target for the Assassin.\n")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.italic = True

        
        insertHR(paragraph)
    
        # Abilities
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
    
        run = paragraph.add_run("Abilities:\n")
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(14)
        font.bold = True
    
        run = paragraph.add_run(get_role_description(player.role) + '\n')
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(11)

        # Secret Information
        if len(player.info) != 0:
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = 0
            paragraph_format.space_after = 0

            run = paragraph.add_run("Secret Information:\n")
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.bold = True
            run = paragraph.add_run('\n'.join(player.info))
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            
        insertHR(paragraph)

        # Cards You Can Play
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
        run = paragraph.add_run("\nPlayable Cards:\n")
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(14)
        font.bold = True
        font.underline = True
        run = paragraph.add_run(get_playable_cards(player.role))
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(12)
        font.bold = True
    
        # Conditional Cards
        first_cards = get_conditional_cards('1-' + player.role)
        second_cards = get_conditional_cards('2-' + player.role)
        third_cards = get_conditional_cards('3-' + player.role)
        fourth_cards = get_conditional_cards('4-' + player.role)
        if first_cards != 'NONE':
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = 0
            paragraph_format.space_after = 0
            run = paragraph.add_run('Conditional Cards: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.bold = True
            font.underline = True
            run = paragraph.add_run('(When you can play them.)\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.bold = True
            font.italic = True
            run = paragraph.add_run(first_cards)
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(12)
            run = paragraph.add_run(second_cards)
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(12)
            font.italic = True
            if third_cards != 'NONE':
                run = paragraph.add_run(third_cards)
                font = run.font
                font.name = 'Caladea'
                font.size = Pt(12)
                font.bold = True
                font.italic = True
                run = paragraph.add_run(fourth_cards)
                font = run.font
                font.name = 'Caladea'
                font.size = Pt(12)
                font.italic = True
        elif third_cards != 'NONE':
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = 0
            paragraph_format.space_after = 0
            run = paragraph.add_run('Conditional Cards: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.bold = True
            font.underline = True
            run = paragraph.add_run('(When you can play them.)\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(14)
            font.bold = True
            font.italic = True
            run = paragraph.add_run(third_cards)
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(12)
            font.bold = True
            font.italic = True
            run = paragraph.add_run(fourth_cards)
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(12)
            font.italic = True
        
        insertHR(paragraph)
    
        # Victory Points
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
    
    
        run = paragraph.add_run("Victory Points:\n")
        font = run.font
        font.name = 'Caladea'
        font.size = Pt(14)
        font.bold = True
    
        # Team Points
        if player.team == 'Good':
            run = paragraph.add_run('Good: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            font.bold = True
            run = paragraph.add_run('If three quests succeeded, you gain 3 Victory Points.\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
        elif player.team == 'Evil':
            run = paragraph.add_run('Evil: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            font.bold = True
            run = paragraph.add_run('If three quests failed, you gain 3 Victory Points.\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
        
        # Origin Points
        if player.origin == 'Mortal' and player.role != 'Sir Kay':
            run = paragraph.add_run('Mortal: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            font.bold = True
            run = paragraph.add_run('If the Holy Grail has become corrupted, you lose 2 Victory Points.\n             If you identified a Fae during The Wild Hunt, you gain 2 Victory Points.\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
        elif player.origin == 'Fae':
            run = paragraph.add_run('Fae: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            font.bold = True
            run = paragraph.add_run('If the Holy Grail has become Corrupted, you gain 2 Victory Points.\n             For each Mortal that identified you during the Wild Hunt, you lose 1 Victory Point.\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
        
        # Assassin Point
        if player.is_assassin == True:
            run = paragraph.add_run('Assassin: ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            font.bold = True
            run = paragraph.add_run('If you correctly assassinated a Target, you gain 1 Victory Point.\n')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
        
        # Role Points
        role_points = get_role_victory_points(player.role)
        if role_points != "NONE":
            run = paragraph.add_run(player.role + ': ')
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            font.bold = True
            run = paragraph.add_run(role_points)
            font = run.font
            font.name = 'Caladea'
            font.size = Pt(11)
            
        player_file = "game/{}.docx".format(player.name)
        document.save(player_file)


    first_player = random.sample(players,1)[0]
    fae_count = 0
    with open("game/1. Read to Start.docx", "w") as file:
        file.write("FIRST LEADER:\nThe player proposing the first mission is {}.\n\n".format(first_player.name))
        file.write("THE HOLY GRAIL\n")
        for player in players:
            if player.origin == "Fae":
                fae_count += 1
            if player.role == "Sir Gawain":
                fae_count -= 1
        if fae_count < 1:
            fae_count = 1
        file.write(f'The Holy Grail is {(fae_count*2) + random.choice([4,5,5])} Fae Spells away from Corruption.')
        #file.write("\n" + second_mission_starter + " is the starting player of the 2nd round.\n")

    with open("game/Do NOT Open.docx", "w") as file:
        file.write("Player -> Role\n\n GOOD TEAM:\n")
        for gp in good_players:
            file.write("{} -> {}\n".format(gp.name, gp.role))
        file.write("\nEVIL TEAM:\n")
        for ep in evil_players:
            file.write("{} -> {}\n".format(ep.name,ep.role))
        if len(neutral_players) > 0:
            file.write("\nNEUTRAL TEAM:\n")
            for np in neutral_players:
                file.write("{} -> {}\n".format(np.name,np.role))
        file.write(f'\nEXCALIBUR:\n')
        for relic in relics:
            if relic == "Excalibur":
                file.write(f'{relic.location}')

if __name__ == "__main__":
#    if not (6 <= len(sys.argv) <= 13):
 #       print("Invalid number of players")
  #      exit(1)

    #players = sys.argv[1:]
    players = ["Jared", "Beka", "Ethan", "Ellie", "Abbie", "Gennie", "Dustin", "Raechel", "Mom", "Hunter", "Marlee", "Clayton"]
    num_players = len(players)
    players = set(players) # use as a set to avoid duplicate players
    players = list(players) # convert to list
    random.shuffle(players) # ensure random order, though set should already do that
    if len(players) != num_players:
        print("No duplicate player names")
        exit(1)

    get_player_info(players)
    