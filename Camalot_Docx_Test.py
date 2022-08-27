from docx import Document
from docx.shared import Pt, RGBColor, Length, Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

document = Document()
sections = document.sections
section = sections[0]

section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0


run = paragraph.add_run("     Player,")
font = run.font
font.name = 'Breathe Fire III'
font.size = Pt(50)

paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0

run = paragraph.add_run("                            you are ")
font = run.font
font.name = 'Breathe Fire III'
font.size = Pt(30)

run = paragraph.add_run("Mordred\n")
font = run.font
font.name = 'Breathe Fire III'
font.size = Pt(30)
font.color.rgb = RGBColor(255, 0, 0)

insertHR(paragraph)

# Allegiance
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0

run = paragraph.add_run("\n     Allegiance: ")
font = run.font
font.name = 'Caladea'
font.size = Pt(14)
font.bold = True

run = paragraph.add_run("Evil")
font = run.font
font.name = 'Caladea'
font.size = Pt(14)

# Origins
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0

run = paragraph.add_run("     Origins: ")
font = run.font
font.name = 'Caladea'
font.size = Pt(14)
font.bold = True

run = paragraph.add_run("Mortal")
font = run.font
font.name = 'Caladea'
font.size = Pt(14)

# Assassination
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0

run = paragraph.add_run("     Assassination Target? ")
font = run.font
font.name = 'Caladea'
font.size = Pt(14)
font.bold = True

run = paragraph.add_run("No\n")
font = run.font
font.name = 'Caladea'
font.size = Pt(14)

insertHR(paragraph)


document.save('word.docx')